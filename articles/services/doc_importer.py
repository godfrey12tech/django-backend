import os
from uuid import uuid4
from django.utils.text import slugify
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from docx import Document
import mammoth

from django.contrib.auth import get_user_model
from articles.models import Article, ArticleImage

User = get_user_model()

class ArticleImporter:
    """
    Parses a .docx file and creates an Article with title, slug, excerpt, content, and embedded images.
    """
    
    def __init__(self, docx_path, author=None):
        if not docx_path:
            raise ValueError("Must provide docx_path")
        self.docx_path = docx_path
        self.author = author or User.objects.filter(role='admin').first()
        if not self.author:
            raise ValueError("No valid author found for import")

    def _load_document(self):
        return Document(self.docx_path)

    def _extract_title(self, doc):
        for paragraph in doc.paragraphs:
            if paragraph.style.name in ("Title", "Heading 1") and paragraph.text.strip():
                return paragraph.text.strip()
        return os.path.basename(self.docx_path).rsplit('.', 1)[0]

    def _extract_excerpt(self, doc, title):
        seen_title = False
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not seen_title and text == title:
                seen_title = True
                continue
            if seen_title and paragraph.style.name == 'Normal' and text:
                return text[:300]
        return ''

    def _convert_body_to_html(self):
        with open(self.docx_path, 'rb') as f:
            result = mammoth.convert_to_html(f)
        return result.value

    def _save_images(self, doc, article):
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                img_data = rel.target_part.blob
                content_type = rel.target_part.content_type
                ext = content_type.split('/')[-1]
                filename = f"article_images/{slugify(article.slug)}_{uuid4().hex}.{ext}"
                path = default_storage.save(filename, ContentFile(img_data))
                ArticleImage.objects.create(article=article, image=path)

    def import_article(self):
        doc = self._load_document()
        title = self._extract_title(doc)
        raw_slug = slugify(title)
        slug = raw_slug[:50]
        excerpt = self._extract_excerpt(doc, title)
        content_html = self._convert_body_to_html()

        article = Article.objects.create(
            title=title,
            slug=slug,
            excerpt=excerpt,
            content=content_html,
            author=self.author,
            # Manual: category, tags, SEO, etc.
        )

        self._save_images(doc, article)
        return article
